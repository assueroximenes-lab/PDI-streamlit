import streamlit as st
import psycopg2
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import altair as alt
from docx import Document
from streamlit_oauth import OAuth2Component
import jwt
from sqlalchemy import create_engine, text
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from datetime import datetime
from io import BytesIO
import os
#from openai import OpenAI


alt.data_transformers.disable_max_rows()

sns.set_theme(style="whitegrid", context="paper")

st.set_page_config(layout="wide")

st.markdown("""
<style>
table { font-size: 11px !important; }
thead tr th { font-size: 11px !important; }
tbody tr td { font-size: 11px !important; }
</style>
""", unsafe_allow_html=True)

# =====================================================
# CONFIG OAUTH GOOGLE
# =====================================================

CLIENT_ID = st.secrets["GOOGLE_CLIENT_ID"]
CLIENT_SECRET = st.secrets["GOOGLE_CLIENT_SECRET"]
REDIRECT_URI = st.secrets["REDIRECT_URI"]

AUTHORIZE_URL = "https://accounts.google.com/o/oauth2/v2/auth"
TOKEN_URL = "https://oauth2.googleapis.com/token"
REFRESH_TOKEN_URL = "https://oauth2.googleapis.com/token"
REVOKE_TOKEN_URL = "https://oauth2.googleapis.com/revoke"

oauth2 = OAuth2Component(
    CLIENT_ID,
    CLIENT_SECRET,
    AUTHORIZE_URL,
    TOKEN_URL,
    REFRESH_TOKEN_URL,
    REVOKE_TOKEN_URL,
)

def verificar_email_google(token):
    decoded = jwt.decode(token, options={"verify_signature": False})
    return decoded.get("email")

# ---------------------------
# CONTROLE DE SESSÃO
# ---------------------------

if "logado" not in st.session_state:
    st.session_state.logado = False
    st.session_state.email = None

# ---------------------------
# LOGIN GOOGLE OBRIGATÓRIO
# ---------------------------

if not st.session_state.logado:

    st.title("Painel Gerencial de Metas")

    result = oauth2.authorize_button(
        name="Entrar com Google (UFAPE)",
        redirect_uri=REDIRECT_URI,
        scope="openid email profile",
        key="google_login"
    )

    if result and "token" in result:

        id_token = result["token"]["id_token"]
        decoded = jwt.decode(id_token, options={"verify_signature": False})
        email = decoded.get("email")

        if email and email.endswith("@ufape.edu.br"):
            st.session_state.logado = True
            st.session_state.email = email
            st.rerun()
        else:
            st.error("Acesso permitido apenas para e-mails institucionais @ufape.edu.br")

    st.stop()

# MOSTRAR USUÁRIO LOGADO
st.sidebar.success(f"Logado como: {st.session_state.email}")

if st.sidebar.button("Sair"):
    st.session_state.clear()
    st.rerun()



# --------------------------
# CONEXÃO
# --------------------------

@st.cache_resource
def get_engine():

    # Se estiver na Streamlit Cloud (tem secrets)
    if "DATABASE_URL" in st.secrets:
        return create_engine(st.secrets["DATABASE_URL"])

    # Se estiver local
    return create_engine(
        "postgresql+psycopg2://postgres:240119@localhost:5432/metas_dev"
    )


def carregar_dados():
    engine = get_engine()
    return pd.read_sql("SELECT * FROM metas", engine)

# ----------------------------
# Função segura para converter inteiro
# ----------------------------
def inteiro_seguro(valor):
    try:
        if valor is None:
            return None
        valor = str(valor).strip()
        if valor == "":
            return None
        return int(float(valor))
    except:
        return None


# ----------------------------
# Função de classificação
# ----------------------------
def classificar_execucao2(row):

    hoje = datetime.now().year

    execucao = row["execucao"]
    inicio = inteiro_seguro(row["Inicio"])
    fim = inteiro_seguro(row["Fim"])
    ano_conclusao = inteiro_seguro(row["Ano_Conclusao"])

    if execucao == "NÃO INICIADA" and inicio is not None and hoje >= inicio:
        return "EXECUÇÃO PENDENTE"

    if execucao != "CONCLUÍDA" and fim is not None and hoje > fim:
        return "ATRASADA"

    if execucao == "NÃO INICIADA" and inicio is not None and hoje < inicio:
        return "A EXECUTAR"

    if execucao in ["INICIADA", "EM ANDAMENTO", "AVANÇADA"]:
        if inicio is not None and fim is not None and inicio <= hoje <= fim:
            return "EM EXECUÇÃO"

    if execucao in ["INICIADA", "EM ANDAMENTO", "AVANÇADA"]:
        if inicio is not None and hoje < inicio:
            return "EXECUÇÃO ANTECIPADA"

    if execucao == "CONCLUÍDA" and ano_conclusao is not None:
        if inicio is not None and fim is not None and inicio <= ano_conclusao <= fim:
            return "CUMPRIDA NO PRAZO"

    if execucao == "CONCLUÍDA" and ano_conclusao is not None:
        if inicio is not None and ano_conclusao < inicio:
            return "CUMPRIDA ANTECIPADA"

    if execucao == "CONCLUÍDA" and ano_conclusao is not None:
        if fim is not None and ano_conclusao > fim:
            return "CUMPRIDA COM ATRASO"

    return "NÃO INICIADA"

def atualizar_execucao2():

    engine = get_engine()

    # 1️⃣ Carrega dados
    df_temp = pd.read_sql("SELECT * FROM metas", engine)

    # 2️⃣ Recalcula a coluna
    df_temp["execucao2"] = df_temp.apply(classificar_execucao2, axis=1)

    # 3️⃣ Atualiza no banco
    with engine.begin() as conn:  # begin já faz commit automático
        for _, row in df_temp.iterrows():
            conn.execute(
                text("""
                    UPDATE metas
                    SET execucao2 = :execucao2
                    WHERE "Ordem" = :ordem
                """),
                {
                    "execucao2": row["execucao2"],
                    "ordem": row["Ordem"]
                }
            )

# --------------------------
# GERAR EXCEL
# --------------------------

def gerar_excel(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Metas")
    return output.getvalue()

# --------------------------
# GERAR WORD (ACRESCENTADO)
# --------------------------
def gerar_relatorio_word(
    df,
    responsavel_sel="Todos",
    eixo_sel="Todos",
    situacao_sel="Todos",
    obj_est_sel="Todos",
    obj_esp_sel="Todos"
):


    # ==========================
    # FUNÇÕES AUXILIARES
    # ==========================
    def add_page_number(document):
        section = document.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        run.text = "Página "

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def set_cell_color(cell, color):
        shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
        cell._tc.get_or_add_tcPr().append(shading)

    def formatar_tabela(tabela, destacar_percentual=False):
        for cell in tabela.rows[0].cells:
            set_cell_color(cell, "D9D9D9")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(11)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i, row in enumerate(tabela.rows[1:]):
            for j, cell in enumerate(row.cells):

                if i % 2 == 0:
                    set_cell_color(cell, "F2F2F2")

                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)

                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER

                if destacar_percentual and j >= len(row.cells) - 2:
                    try:
                        valor = float(cell.text.replace(",", "."))
                        if valor >= 70:
                            set_cell_color(cell, "C6EFCE")
                        elif valor < 40:
                            set_cell_color(cell, "FFC7CE")
                    except:
                        pass

    # ==========================
    # PREPARAÇÃO
    # ==========================
    df_rel = df.copy()

    if "Resp_1" in df_rel.columns:
        df_rel["Resp_1"] = df_rel["Resp_1"].astype(str).str.strip()

    document = Document()
    add_page_number(document)

    # ==========================
    # RESPONSÁVEL AUTOMÁTICO
    # ==========================
    if "Resp_1" in df_rel.columns and not df_rel.empty:
        r = df_rel["Resp_1"].dropna().unique()
        responsavel_final = r[0] if len(r) == 1 else "Geral"
    else:
        responsavel_final = "Geral"

    # ==========================
    # CAPA INSTITUCIONAL
    # ==========================
    try:
        document.add_picture("logo_ufape.png", width=Inches(2))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        pass

    p = document.add_paragraph()
    p.add_run("\nUNIVERSIDADE FEDERAL DO AGRESTE DE PERNAMBUCO\n").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph()
    p.add_run("RELATÓRIO GERENCIAL DO PDI\n").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph()
    p.add_run(f"Responsável: {responsavel_final}\n")
    p.add_run(f"Data e hora: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    # ==========================
    # ESCOPO
    # ==========================
    document.add_heading("Escopo do Relatório", 1)

    filtros = {
        "Responsável": responsavel_final,
        "Situação": situacao_sel if situacao_sel != "Todos" else "Todas",
        "Eixo": eixo_sel if eixo_sel != "Todos" else "Todos",
        "Objetivo Estratégico": obj_est_sel if obj_est_sel != "Todos" else "Todos",
        "Objetivo Específico": obj_esp_sel if obj_esp_sel != "Todos" else "Todos",
    }

    for k, v in filtros.items():
        p = document.add_paragraph()
        p.add_run(f"{k}: ").bold = True
        p.add_run(str(v))

    # ==========================
    # 1. RESUMO GERAL
    # ==========================
    document.add_heading("1. Resumo Geral", 1)

    resumo = df_rel.groupby("execucao").size().reset_index(name="Quantidade")
    total = resumo["Quantidade"].sum()
    resumo["Percentual (%)"] = (resumo["Quantidade"] / total * 100).round(1) if total else 0

    tabela = document.add_table(rows=1, cols=3)
    tabela.style = "Table Grid"

    cab = tabela.rows[0].cells
    cab[0].text = "Situação"
    cab[1].text = "Quantidade"
    cab[2].text = "Percentual (%)"

    for _, row in resumo.iterrows():
        l = tabela.add_row().cells
        l[0].text = str(row["execucao"])
        l[1].text = str(row["Quantidade"])
        l[2].text = str(row["Percentual (%)"])

    formatar_tabela(tabela, True)

    # ==========================
    # 2. EXECUÇÃO POR EIXO (COMPLETO)
    # ==========================
    if "Eixo" in df_rel.columns:

        document.add_page_break()
        document.add_heading("2. Execução por Eixo", 1)

        # criar status binário (igual aos outros itens)
        df_rel["status"] = df_rel["execucao"].apply(
            lambda x: "Concluída" if "conclu" in str(x).lower() else "Não Concluída"
        )

        resumo_eixo = (
            df_rel.groupby(["Eixo", "status"])
            .size()
            .unstack(fill_value=0)
            .reset_index()
        )

        # garantir colunas
        if "Concluída" not in resumo_eixo.columns:
            resumo_eixo["Concluída"] = 0
        if "Não Concluída" not in resumo_eixo.columns:
            resumo_eixo["Não Concluída"] = 0

        # totais
        resumo_eixo["Total"] = resumo_eixo["Concluída"] + resumo_eixo["Não Concluída"]

        # percentuais
        resumo_eixo["% Concluída"] = (
                resumo_eixo["Concluída"] / resumo_eixo["Total"] * 100
        ).round(1)

        resumo_eixo["% Não Concluída"] = (
                resumo_eixo["Não Concluída"] / resumo_eixo["Total"] * 100
        ).round(1)

        # ordenar
        resumo_eixo = resumo_eixo.sort_values("Total", ascending=False)

        # tabela
        tabela = document.add_table(rows=1, cols=6)
        tabela.style = "Table Grid"

        cab = tabela.rows[0].cells
        cab[0].text = "Eixo"
        cab[1].text = "Total"
        cab[2].text = "Concluídas"
        cab[3].text = "Não Concluídas"
        cab[4].text = "% Concluída"
        cab[5].text = "% Não Concluída"

        for _, row in resumo_eixo.iterrows():
            l = tabela.add_row().cells
            l[0].text = str(row["Eixo"])
            l[1].text = str(row["Total"])
            l[2].text = str(row["Concluída"])
            l[3].text = str(row["Não Concluída"])
            l[4].text = str(row["% Concluída"])
            l[5].text = str(row["% Não Concluída"])

        formatar_tabela(tabela, True)

    # ==========================
    # 3. RESPONSÁVEIS
    # ==========================
    if "Resp_1" in df_rel.columns:

        document.add_page_break()
        document.add_heading("3. Ranking de Responsáveis", 1)

        df_rel["status"] = df_rel["execucao"].apply(
            lambda x: "Concluída" if "conclu" in str(x).lower() else "Não Concluída"
        )

        r = df_rel.groupby(["Resp_1", "status"]).size().unstack(fill_value=0).reset_index()

        r["Total"] = r.sum(axis=1, numeric_only=True)
        r["% Concluída"] = (r.get("Concluída", 0) / r["Total"] * 100).round(1)
        r["% Não Concluída"] = (r.get("Não Concluída", 0) / r["Total"] * 100).round(1)

        tabela = document.add_table(rows=1, cols=6)
        tabela.style = "Table Grid"

        cab = tabela.rows[0].cells
        cab[0].text = "Responsável"
        cab[1].text = "Total"
        cab[2].text = "Concluídas"
        cab[3].text = "Não Concluídas"
        cab[4].text = "% Concluída"
        cab[5].text = "% Não Concluída"

        for _, row in r.iterrows():
            l = tabela.add_row().cells
            l[0].text = str(row["Resp_1"])
            l[1].text = str(row["Total"])
            l[2].text = str(row.get("Concluída", 0))
            l[3].text = str(row.get("Não Concluída", 0))
            l[4].text = str(row["% Concluída"])
            l[5].text = str(row["% Não Concluída"])

        formatar_tabela(tabela, True)

    # ==========================
    # 4. OBJETIVO ESTRATÉGICO
    # ==========================
    if "Objetivo Estratégico" in df_rel.columns:

        document.add_page_break()
        document.add_heading("4. Análise por Objetivo Estratégico", 1)

        o = df_rel.groupby(["Objetivo Estratégico", "status"]).size().unstack(fill_value=0).reset_index()

        o["Total"] = o.sum(axis=1, numeric_only=True)
        o["% Concluída"] = (o.get("Concluída", 0) / o["Total"] * 100).round(1)
        o["% Não Concluída"] = (o.get("Não Concluída", 0) / o["Total"] * 100).round(1)

        tabela = document.add_table(rows=1, cols=6)
        tabela.style = "Table Grid"

        cab = tabela.rows[0].cells
        cab[0].text = "Objetivo Estratégico"
        cab[1].text = "Total"
        cab[2].text = "Concluídas"
        cab[3].text = "Não Concluídas"
        cab[4].text = "% Concluída"
        cab[5].text = "% Não Concluída"

        for _, row in o.iterrows():
            l = tabela.add_row().cells
            l[0].text = str(row["Objetivo Estratégico"])
            l[1].text = str(row["Total"])
            l[2].text = str(row.get("Concluída", 0))
            l[3].text = str(row.get("Não Concluída", 0))
            l[4].text = str(row["% Concluída"])
            l[5].text = str(row["% Não Concluída"])

        formatar_tabela(tabela, True)

    # ==========================
    # FINAL (OBRIGATÓRIO)
    # ==========================
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()

# --------------------------
# CARREGAR DADOS
# --------------------------
atualizar_execucao2()
df_original = carregar_dados()


if df_original.empty:
    st.warning("Nenhuma meta encontrada no banco.")
    st.stop()

df_original["execucao"] = df_original["execucao"].replace("", None)
df_original["execucao"] = df_original["execucao"].fillna("NÃO INICIADA")

situacoes_padrao = [
    "NÃO INICIADA",
    "INICIADA",
    "EM ANDAMENTO",
    "AVANÇADA",
    "CONCLUÍDA"
]

paleta = {
    "NÃO INICIADA": "#9aa0a6",
    "INICIADA": "#4e79a7",
    "EM ANDAMENTO": "#f28e2b",
    "AVANÇADA": "#59a14f",
    "CONCLUÍDA": "#2ca02c"
}

situacoes_exec2 = [
    "A EXECUTAR",
    "EXECUÇÃO PENDENTE",
    "EM EXECUÇÃO",
    "EXECUÇÃO ANTECIPADA",
    "ATRASADA",
    "CUMPRIDA NO PRAZO",
    "CUMPRIDA ANTECIPADA",
    "CUMPRIDA COM ATRASO",
    "NÃO INICIADA"
]

paleta_exec2 = {
    "A EXECUTAR": "#4e79a7",
    "EXECUÇÃO PENDENTE": "#f28e2b",
    "EM EXECUÇÃO": "#59a14f",
    "EXECUÇÃO ANTECIPADA": "#76b7b2",
    "ATRASADA": "#e15759",
    "CUMPRIDA NO PRAZO": "#2ca02c",
    "CUMPRIDA ANTECIPADA": "#8cd17d",
    "CUMPRIDA COM ATRASO": "#ff9da7",
    "NÃO INICIADA": "#9aa0a6"
}

# -----------------------------------------------------------------------
# FILTROS
# -----------------------------------------------------------------------

st.sidebar.title("Filtros")

# ==========================
# RESPONSÁVEL
# ==========================
responsaveis = ["Todos"] + sorted(
    df_original["Resp_1"].dropna().unique()
)
responsavel_sel = st.sidebar.selectbox("Responsável", responsaveis)

# ==========================
# SITUAÇÃO
# ==========================
situacoes = ["Todos"] + situacoes_padrao
situacao_sel = st.sidebar.selectbox("Situação", situacoes)

# ==========================
# EIXO
# ==========================
eixos = ["Todos"] + sorted(
    df_original["Eixo"].dropna().unique()
)
eixo_sel = st.sidebar.selectbox("Eixo", eixos)

# Base intermediária para dependência
df_temp = df_original.copy()

if eixo_sel != "Todos":
    df_temp = df_temp[df_temp["Eixo"] == eixo_sel]

# ==========================
# OBJETIVO ESTRATÉGICO
# ==========================
obj_est = ["Todos"] + sorted(
    df_temp["Objetivo Estratégico"].dropna().unique()
)
obj_est_sel = st.sidebar.selectbox("Objetivo Estratégico", obj_est)

if obj_est_sel != "Todos":
    df_temp = df_temp[df_temp["Objetivo Estratégico"] == obj_est_sel]

# ==========================
# OBJETIVO ESPECÍFICO
# ==========================
obj_esp = ["Todos"] + sorted(
    df_temp["Objetivo Específico"].dropna().unique()
)
obj_esp_sel = st.sidebar.selectbox("Objetivo Específico", obj_esp)


# --------------------------
# APLICAR FILTROS FINAIS
# --------------------------

df = df_original.copy()

if responsavel_sel != "Todos":
    df = df[df["Resp_1"] == responsavel_sel]

if situacao_sel != "Todos":
    df = df[df["execucao"] == situacao_sel]

if eixo_sel != "Todos":
    df = df[df["Eixo"] == eixo_sel]

if obj_est_sel != "Todos":
    df = df[df["Objetivo Estratégico"] == obj_est_sel]

if obj_esp_sel != "Todos":
    df = df[df["Objetivo Específico"] == obj_esp_sel]


# --------------------------
# DOWNLOAD EXCEL
# --------------------------

st.sidebar.markdown("---")
st.sidebar.subheader("Exportação")

excel_bytes = gerar_excel(df_original)

st.sidebar.download_button(
    label="Baixar base em Excel",
    data=excel_bytes,
    file_name="metas.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)

# --------------------------
# BOTÃO WORD (ACRESCENTADO)
# --------------------------

#word_bytes = gerar_relatorio_word(df_original)
word_bytes = gerar_relatorio_word(df)
st.sidebar.download_button(
    label="Baixar relatório em Word",
    data=word_bytes,
    file_name="relatorio_metas.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)



# --------------------------
# DASHBOARD
# --------------------------

st.title("Painel Gerencial de Metas")

total_metas = len(df)
st.metric("Total de Metas", total_metas)

st.divider()


# --------------------------
# SITUAÇÃO DAS METAS
# --------------------------

dados = []
for s in situacoes_padrao:
    qtd = df[df["execucao"] == s].shape[0]
    perc = (qtd / total_metas * 100) if total_metas else 0
    dados.append([s, qtd, round(perc,1)])

tabela_situacao = pd.DataFrame(
    dados,
    columns=["Situação", "Quantidade", "Percentual (%)"]
)

# --------------------------
# KPIs EM CIMA (HORIZONTAL)
# --------------------------

st.markdown("""
<style>

/* NÚMERO GRANDE (Quantidade) */
div[data-testid="stMetricValue"] {
    font-size: 48px !important;   /* 👈 3x maior */
    font-weight: 800;
}

/* PERCENTUAL (delta) */
div[data-testid="stMetricDelta"] {
    font-size: 22px !important;   /* 👈 maior e legível */
    font-weight: 600;
}

/* TÍTULO DO KPI (Situação) */
div[data-testid="stMetricLabel"] {
    font-size: 16px;
    font-weight: 600;
}

</style>
""", unsafe_allow_html=True)

st.markdown("### 📊 Situação das Metas")

cols = st.columns(len(tabela_situacao))

for i, row in tabela_situacao.iterrows():
    cols[i].metric(
        label=row["Situação"],
        value=row["Quantidade"],
        delta=f'{row["Percentual (%)"]}%'
    )

st.divider()

# --------------------------
# GRÁFICO GRANDE EMBAIXO (ALTair)
# --------------------------

df_pizza = tabela_situacao[tabela_situacao["Quantidade"] > 0]

chart = alt.Chart(df_pizza).mark_arc(innerRadius=90).encode(
    theta=alt.Theta(field="Quantidade", type="quantitative"),
    color=alt.Color(
        field="Situação",
        type="nominal",
        legend=alt.Legend(title="Situação", orient="right")
    ),
    tooltip=["Situação", "Quantidade", "Percentual (%)"]
).properties(
    width=600,   # 👈 4x maior (antes ~150)
    height=400,
    title="Distribuição das Metas"
)

# centralizar
col1, col2, col3 = st.columns([1,2,1])
with col2:
    st.altair_chart(chart, width="stretch")

st.divider()



# --------------------------
# GRÁFICO SITUAÇÃO
# --------------------------

colA, colB = st.columns([1,1])

with colA:
    st.dataframe(tabela_situacao, width="stretch")

with colB:

    grafico1 = (
        alt.Chart(tabela_situacao)
        .mark_bar()
        .encode(
            y=alt.Y("Situação:N", sort="-x", title=None),
            x=alt.X("Quantidade:Q", title=None),
            color=alt.Color(
                "Situação:N",
                scale=alt.Scale(
                    domain=situacoes_padrao,
                    range=[paleta[s] for s in situacoes_padrao]
                ),
                legend=None
            ),
            tooltip=["Situação", "Quantidade", "Percentual (%)"]
        )
        .properties(height=400, width=700)
    )

    st.altair_chart(grafico1)

st.divider()

# --------------------------
# EXECUÇÃO POR RESPONSÁVEL
# --------------------------

st.subheader("Execução por Responsável")

tabela_resp = (
    df.groupby(["Resp_1", "execucao"])
    .size()
    .reset_index(name="Quantidade")
)

altura = max(200, 25 * tabela_resp["Resp_1"].nunique())

grafico = (
    alt.Chart(tabela_resp)
    .mark_bar()
    .encode(
        y=alt.Y(
            "Resp_1:N",
            sort="-x",
            title=None,
            axis=alt.Axis(labelLimit=1000, labelOverlap=False)
        ),
        x=alt.X("Quantidade:Q", title=None),
        color=alt.Color(
            "execucao:N",
            scale=alt.Scale(
                domain=situacoes_padrao,
                range=[paleta[s] for s in situacoes_padrao]
            ),
            legend=alt.Legend(orient="bottom", columns=3)
        ),
        tooltip=["Resp_1", "execucao", "Quantidade"]
    )
    .properties(height=altura)
)

st.altair_chart(grafico, width="stretch")

st.divider()

# --------------------------
# RESUMO POR EIXO E SITUAÇÃO
# --------------------------

st.subheader("Resumo de Metas por Eixo e Situação")

if "Eixo" not in df.columns:
    st.warning("A coluna 'Eixo' não existe na tabela metas.")
else:

    tabela_eixo = pd.pivot_table(
        df,
        index="Eixo",
        columns="execucao",
        values="Meta",
        aggfunc="count",
        fill_value=0
    )

    for s in situacoes_padrao:
        if s not in tabela_eixo.columns:
            tabela_eixo[s] = 0

    tabela_eixo = tabela_eixo[situacoes_padrao]
    tabela_eixo["TOTAL"] = tabela_eixo.sum(axis=1)

    st.dataframe(tabela_eixo, width="stretch")

# ---------------------------------------------
# METAS DETALHADAS GERAIS
# ---------------------------------------------

st.subheader("Metas Detalhadas (Visão Geral)")

colunas_exibir = ["Eixo", "Resp_1", "Meta", "Descrição da Meta", "execucao"]
colunas_existentes = [c for c in colunas_exibir if c in df.columns]

st.dataframe(df[colunas_existentes],width="stretch")

st.divider()
st.subheader("Situação Temporal por Responsável")
total_temporal = len(df)
st.metric("Total de Metas (Temporal)", total_temporal)

if "execucao2" in df.columns:

    total_temporal = len(df)

    dados_exec2 = []
    for s in situacoes_exec2:
        qtd = df[df["execucao2"] == s].shape[0]
        perc = (qtd / total_temporal * 100) if total_temporal else 0
        dados_exec2.append([s, qtd, round(perc, 1)])

    tabela_indicadores_exec2 = pd.DataFrame(
        dados_exec2,
        columns=["Situação", "Quantidade", "Percentual (%)"]
    )

    # cria colunas dinâmicas (3 por linha para não ficar esmagado)
    n_colunas = 3
    for i in range(0, len(tabela_indicadores_exec2), n_colunas):
        cols = st.columns(n_colunas)
        for j in range(n_colunas):
            if i + j < len(tabela_indicadores_exec2):
                row = tabela_indicadores_exec2.iloc[i + j]
                cols[j].metric(
                    label=row["Situação"],
                    value=row["Quantidade"],
                    delta=f'{row["Percentual (%)"]}%'
                )

    st.divider()



if "execucao2" in df.columns:

    tabela_exec2 = (
        df.groupby(["Resp_1", "execucao2"])
        .size()
        .reset_index(name="Quantidade")
    )

    altura2 = max(200, 25 * tabela_exec2["Resp_1"].nunique())

    grafico_exec2 = (
        alt.Chart(tabela_exec2)
        .mark_bar()
        .encode(
            y=alt.Y(
                "Resp_1:N",
                sort="-x",
                title=None,
                axis=alt.Axis(labelLimit=1000, labelOverlap=False)
            ),
            x=alt.X("Quantidade:Q", title=None),
            color=alt.Color(
                "execucao2:N",
                scale=alt.Scale(
                    domain=situacoes_exec2,
                    range=[paleta_exec2[s] for s in situacoes_exec2]
                ),
                legend=alt.Legend(orient="bottom", columns=3)
            ),
            tooltip=["Resp_1", "execucao2", "Quantidade"]
        )
        .properties(height=altura2)
    )

    st.altair_chart(grafico_exec2,width="stretch")

else:
    st.warning("A coluna execucao2 ainda não foi gerada.")



st.divider()


st.divider()
st.subheader("Situação Temporal por Responsável")

if "execucao2" in df.columns:
    tabela_exec2 = (
        df.groupby(["Resp_1", "execucao2"])
        .size()
        .reset_index(name="Quantidade")
    )
    st.dataframe(tabela_exec2, width="stretch")